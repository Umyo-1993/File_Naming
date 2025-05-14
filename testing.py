import os
from datetime import datetime
from docx import Document  # Import the docx library

commit = input("Please enter your commit message: ").strip()
version = input("Please enter your version of the file here: ").strip()

now = datetime.now()
timestamp = now.strftime("%Y-%m-%d_%H-%M-%S")

safe_commit = "".join(c for c in commit if c.isalnum() or c in (" ", "_", "-")).rstrip().replace(" ", "_")
safe_version = "".join(c for c in version if c.isalnum() or c in ("_", "-", "."))

# Change the extension to .docx
file_name = f"{safe_commit}_{timestamp}_v{safe_version}.docx"

# Create a new Document object
document = Document()

# Add a title to the document
document.add_heading("Commit Information", level=1)

# Add the commit message and version as paragraphs
document.add_paragraph(f"Commit Message: {commit}")
document.add_paragraph(f"Version: {version}")
document.add_paragraph(f"Timestamp: {timestamp}")

# Save the document
document.save(file_name)

print("✅ Created:", file_name)
print("Make sure you have the python-docx library installed (`pip install python-docx`)")
